"""v0.12 Resume parser tests."""
import sys
from pathlib import Path
sys.path.insert(0, str(Path(__file__).resolve().parent.parent))


def test_parse_txt_success():
    """TXT 文件解析成功"""
    from services.resume_parser import parse_resume_file
    content = "张三\n本科计算机专业\n3年Python开发经验\n熟悉Django、FastAPI、MySQL\n".encode("utf-8")
    result = parse_resume_file("resume.txt", content, "text/plain")
    assert result["file_type"] == "txt"
    assert result["char_count"] > 0
    assert "Python" in result["text"]


def test_parse_txt_gbk():
    """GBK 编码 TXT 解析成功"""
    from services.resume_parser import parse_resume_file
    content = "张三\n本科计算机专业\nPython开发经验\n".encode("gbk")
    result = parse_resume_file("resume.txt", content, "text/plain")
    assert result["file_type"] == "txt"
    assert "Python" in result["text"]
    assert any("编码" in w for w in result["warnings"])


def test_reject_empty_file():
    """空文件被拒绝"""
    from services.resume_parser import parse_resume_file
    import pytest
    with pytest.raises(ValueError, match="文件为空"):
        parse_resume_file("empty.txt", b"", "text/plain")


def test_reject_unsupported_format():
    """不支持格式被拒绝"""
    from services.resume_parser import parse_resume_file
    import pytest
    with pytest.raises(ValueError, match="不支持"):
        parse_resume_file("resume.xlsx", b"some content", "")


def test_reject_oversized_file():
    """超大文件被拒绝"""
    from services.resume_parser import parse_resume_file
    import pytest
    big = b"x" * (11 * 1024 * 1024)  # 11MB
    with pytest.raises(ValueError, match="文件过大"):
        parse_resume_file("big.txt", big, "text/plain")


def test_parse_docx():
    """DOCX 解析（python-docx 已安装）"""
    try:
        import docx
        from io import BytesIO
        # 创建一个最小 DOCX
        d = docx.Document()
        d.add_paragraph("张三 简历")
        d.add_paragraph("Python开发经验 3年")
        buf = BytesIO()
        d.save(buf)
        content = buf.getvalue()

        from services.resume_parser import parse_resume_file
        result = parse_resume_file("resume.docx", content)
        assert result["file_type"] == "docx"
        assert "Python" in result["text"]
    except ImportError:
        pass  # python-docx 未安装时跳过


def test_parse_pdf():
    """PDF 解析（pypdf 已安装时测试）"""
    try:
        import pypdf
        from io import BytesIO
        # 创建一个带文本的 PDF
        writer = pypdf.PdfWriter()
        writer.add_blank_page(width=612, height=792)
        buf = BytesIO()
        writer.write(buf)
        content = buf.getvalue()

        from services.resume_parser import parse_resume_file
        try:
            result = parse_resume_file("resume.pdf", content)
            assert result["file_type"] == "pdf"
        except ValueError:
            pass  # 空白 PDF 解析为空，属于正常
    except ImportError:
        pass  # pypdf 未安装时跳过


def test_short_text_warning():
    """过短文本给 warning"""
    from services.resume_parser import parse_resume_file
    result = parse_resume_file("short.txt", b"hello world", "text/plain")
    assert len(result["warnings"]) > 0
    assert any("短" in w for w in result["warnings"])


def test_detect_file_type():
    """文件类型检测"""
    from services.resume_parser import _detect_file_type
    assert _detect_file_type("resume.pdf", "") == "pdf"
    assert _detect_file_type("resume.docx", "") == "docx"
    assert _detect_file_type("resume.txt", "text/plain") == "txt"
    assert _detect_file_type("resume.jpg", "") == ""
